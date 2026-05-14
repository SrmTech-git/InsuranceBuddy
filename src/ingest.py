# document loading and chunking

import re
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document

from chunkers import get_chunker


def parse_filename(file_path: str) -> dict:
    """Extract form number, edition date, and description from a PDF filename.

    Parsing strategy:
    1. Use regex to find a form number, edition date, and description
       directly from the full path string — avoids os.path.basename which
       misinterprets '/' in dates like '02/2021' as path separators on Windows.
    2. Fall back gracefully if any part is missing.
    """

    form_number = ""
    edition_date = ""
    description = ""
    parsed = True

    # Master pattern that captures all three parts in one pass:
    #   form_number  = 2-6 uppercase letters + optional space + digits
    #                  + optional alphabetic suffix (e.g. ACORD 0050WM)
    #   edition_date = parenthesized group right after the form number
    #   description  = whatever text follows
    pattern = (
        r"([A-Z]{2,6}\s?\d+[A-Z]*(?:[.\-]\d+)*)"   # group 1: form number
        r"\s*\(([^)]+)\)"        # group 2: edition date in parens
        r"\s*(.+?)\.(?:pdf|txt|docx)"    # group 3: description before extension
    )
    full_match = re.search(pattern, file_path, re.IGNORECASE)

    if full_match:
        form_number = full_match.group(1).strip()
        edition_date = full_match.group(2).strip()
        description = full_match.group(3).strip(" -–—")
    else:
        # Try matching just a form number (no edition date)
        partial = re.search(
            r"([A-Z]{2,6}\s?\d+[A-Z]*(?:[.\-]\d+)*)\s*(.+?)\.(?:pdf|txt|docx)",
            file_path, re.IGNORECASE,
        )
        if partial:
            form_number = partial.group(1).strip()
            remainder = partial.group(2).strip(" -–—")
            # Check if remainder starts with a parenthesized group
            date_match = re.match(r"\(([^)]+)\)\s*(.*)", remainder)
            if date_match:
                edition_date = date_match.group(1).strip()
                description = date_match.group(2).strip(" -–—")
            else:
                description = remainder
        else:
            # No form number found at all — pull description from the filename
            # Strip directory prefixes and extension manually
            name = re.sub(r"\.(?:pdf|txt|docx)$", "", file_path, flags=re.IGNORECASE)
            # Remove any directory path (split on / and \ )
            name = re.split(r"[\\/]", name)[-1]
            description = name.strip()
            parsed = False

    # If we got a form number but nothing else, flag as partially parsed
    if form_number and not edition_date and not description:
        parsed = False

    # Build display filename by stripping directory prefixes.
    # We strip on '\' always, and strip on '/' only when the slash
    # appears before any parenthesized group (real dir separator,
    # not a date like '02/2021' inside parens).
    display_name = file_path.rsplit("\\", 1)[-1]
    # Find the position of the first '(' — slashes before it are directories
    paren_pos = display_name.find("(")
    if paren_pos == -1:
        # No parens at all — safe to split on '/' normally
        display_name = display_name.rsplit("/", 1)[-1]
    else:
        # Only strip directory slashes that appear before the first '('
        prefix = display_name[:paren_pos]
        if "/" in prefix:
            dir_end = prefix.rfind("/") + 1
            display_name = display_name[dir_end:]

    metadata = {
        "form_number": form_number,
        "edition_date": edition_date,
        "description": description,
        "filename": display_name,
        "parsed": parsed,
    }

    return metadata


def _load_docx(file_path: str) -> list[Document]:
    """Extract text from a .docx file and return a single LangChain Document."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    text = "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
    return [Document(page_content=text, metadata={"source": file_path})]


def load_and_split(file_path: str, collection_name: str = "") -> list:
    """Load a document and split it into chunks with parsed filename metadata.

    Supports .pdf, .txt, and .docx files. For .xlsx, use
    `ingest_xlsx.load_xlsx_by_state()` directly — its per-state output
    shape doesn't fit this function's flat list[Document] return type.

    The chunker is selected by collection_name via the chunkers registry.
    Unregistered collections (or empty collection_name) fall back to the
    default character-based chunker.
    """

    # Parse metadata from the filename
    file_metadata = parse_filename(file_path)

    print("Parsed metadata:")
    for key, value in file_metadata.items():
        print(f"  {key}: {value}")

    # Pick the right loader based on file extension
    ext = file_path.rsplit(".", 1)[-1].lower() if "." in file_path else ""
    if ext == "pdf":
        loader = PyPDFLoader(file_path)
        pages = loader.load()
    elif ext == "txt":
        loader = TextLoader(file_path, encoding="utf-8")
        pages = loader.load()
    elif ext == "docx":
        pages = _load_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file type: .{ext} — supported: .pdf, .txt, .docx. "
            f"For .xlsx, call ingest_xlsx.load_xlsx_by_state() directly."
        )

    print(f"\nLoaded {len(pages)} page(s) from {file_path}")

    # Look up the chunker for this collection — falls back to default
    # for unregistered collections.
    chunker = get_chunker(collection_name)
    chunks = chunker(pages, file_metadata)

    print(f"Created {len(chunks)} chunks")

    return chunks


if __name__ == "__main__":
    # Smoke test: filename parsing only (no file I/O).
    test_filename = "INS7059 (Rev. 02/2021) Instructions For Filing Annual & Quarterly Statements - MEWAs.pdf"
    print(f"Testing filename parsing on: {test_filename}\n")
    metadata = parse_filename(test_filename)
    for key, value in metadata.items():
        print(f"  {key}: {value}")
